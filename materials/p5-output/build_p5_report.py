from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "materials" / "p5-output"
REPORT_PATH = OUT_DIR / "P5 Comprehensive MVP security analysis.docx"


def set_run_font(run, *, size: float = 14, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def set_cell_margins(cell, margin: int = 90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge in ("top", "start", "bottom", "end"):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(margin))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_cm: list[float]) -> None:
    table.autofit = False
    table.allow_autofit = False
    table.style = "Table Grid"

    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_width = tbl_pr.find(qn("w:tblW"))
    if tbl_width is None:
        tbl_width = OxmlElement("w:tblW")
        tbl_pr.append(tbl_width)
    tbl_width.set(qn("w:type"), "dxa")
    tbl_width.set(qn("w:w"), str(int(sum(widths_cm) / 2.54 * 1440)))

    grid = table._tbl.tblGrid
    for child in list(grid.gridCol_lst):
        grid.remove(child)
    for width in widths_cm:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width / 2.54 * 1440)))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[index])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(widths_cm[index] / 2.54 * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.PORTRAIT
    section.page_width = Cm(27.94)
    section.page_height = Cm(43.18)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)


def add_paragraph(
    doc: Document,
    text: str = "",
    *,
    bold: bool = False,
    italic: bool = False,
    size: float = 14,
    align: int | None = None,
) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=14, bold=True)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        set_run_font(run, size=14)


def add_code(doc: Document, text: str, *, size: float = 9) -> None:
    for line in text.strip("\n").splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(size)


def add_d2_listing(
    doc: Document,
    listing_title: str,
    description: str,
    d2_code: str,
    flows: list[list[str]],
) -> None:
    add_heading(doc, listing_title)
    add_paragraph(doc, description)
    add_code(doc, f"```d2\n{d2_code.strip()}\n```", size=8.2)
    add_paragraph(doc, "Таблица потоков для данной диаграммы:")
    add_table(
        doc,
        ["Flow ID", "Source", "Target", "Data / Action", "Control / Protection"],
        flows,
        [2.0, 5.3, 5.3, 6.7, 7.8],
        font_size=8.2,
    )


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths_cm: list[float],
    *,
    font_size: float = 9.5,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths_cm)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run_font(run, size=font_size, bold=True)

    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cells[index].text = value
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run_font(run, size=font_size)

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
    doc.add_paragraph()


def snippet(relative: str, start: int, end: int) -> str:
    path = ROOT / relative
    lines = path.read_text(encoding="utf-8").splitlines()
    return "\n".join(lines[start - 1 : end])


def annotated_snippet(notes: list[str], relative: str, start: int, end: int) -> str:
    comments = "\n".join(f"# {note}" for note in notes)
    return f"{comments}\n{snippet(relative, start, end)}"


def annotated_multirange_snippet(
    notes: list[str],
    relative: str,
    ranges: list[tuple[int, int]],
) -> str:
    comments = "\n".join(f"# {note}" for note in notes)
    code = "\n\n# ...\n\n".join(snippet(relative, start, end) for start, end in ranges)
    return f"{comments}\n{code}"


def save_with_fallback(doc: Document, output_path: Path) -> Path:
    try:
        doc.save(output_path)
        return output_path
    except PermissionError:
        fallback = output_path.with_name(output_path.stem + " - regenerated" + output_path.suffix)
        doc.save(fallback)
        return fallback


def add_title(doc: Document) -> None:
    lines = [
        ("Практическая работа №5.", True),
        (
            "Тема: Комплексный анализ безопасности MVP: память, ресурсы, обработка данных, "
            "аутентификация, авторизация и криптографическая защита.",
            True,
        ),
        ("Вариант: 2. Oil and gas asset maintenance MVP.", False),
    ]
    for text, bold in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, size=14, bold=bold)

    add_paragraph(doc, "Выполнил: Нурым Абзал, CSE-2506M.")
    add_paragraph(doc, "Проверил: Rostyslav Lisnevskyi")


def add_intro(doc: Document) -> None:
    add_paragraph(
        doc,
        "Цель работы. Выполнить углубленный анализ безопасности MVP, проверить обработку памяти "
        "и ресурсов, точки входа пользовательских данных, dangerous sinks, а также механизмы "
        "authentication, authorization и cryptography. По результатам анализа требуется исправить "
        "реальные уязвимости, подтвердить исправления тестами и рассчитать CVSS v4.0 для пяти "
        "новых уязвимостей, не повторяющих P4.",
        bold=True,
    )
    add_paragraph(
        doc,
        "В качестве базового состояния для P5 использован результат P4: "
        "materials/p4-snapshots/mvp-after. Текущее исправленное состояние находится в папке mvp. "
        "Сравнение снапшота показало ожидаемые изменения только в P5-файлах: конфигурация, "
        "маршруты списков, authorization, authentication service, pyproject.toml и новые тесты.",
    )
    add_paragraph(
        doc,
        "Термины использованы в логике лекций L7 и L8: пользовательский ввод рассматривается как "
        "source, опасная операция как sink, а trust boundary показывает место, где данные переходят "
        "из недоверенной зоны в доверенную. Для L8 отдельно различаются authentication "
        "(установление личности), authorization (право на действие или объект) и cryptography "
        "(хеширование паролей, токенов и защита секретов).",
    )
    add_d2_listing(
        doc,
        "Листинг 1. D2-код структурной схемы MVP и trust boundaries.",
        "Код ниже описывает структурную схему в формате D2. Его можно отрендерить через D2 CLI в SVG и вставить в draw.io или Word. Диаграмма показывает внешнюю недоверенную зону, границу доверия приложения и основные компоненты обработки данных.",
        """
direction: right

external: {
  label: "External / untrusted zone"
  style.fill: "#F4F1EA"
  client: "Browser / API client"
  frontend: "Static frontend /app"
}

application: {
  label: "Application trust boundary"
  style.fill: "#FBFAF7"
  routes: "FastAPI routes"
  dependencies: "Dependencies: current user and DB session"
  validation: "Pydantic validation and bounded query params"
  services: "Services and business rules"
  authorization: "Authorization: role and object matrix"
  audit: "Audit log helper"
  orm: "SQLAlchemy ORM"
  database: {
    label: "SQLite / production DB"
    shape: cylinder
  }
  refresh_store: {
    label: "Refresh token hash store"
    shape: cylinder
  }
}

external.client -> application.routes: "JSON, form data, query params, headers"
external.frontend -> application.routes: "same-origin fetch; textContent rendering"
application.routes -> application.validation
application.routes -> application.dependencies
application.dependencies -> application.authorization
application.validation -> application.services
application.authorization -> application.services
application.services -> application.orm
application.orm -> application.database
application.services -> application.refresh_store
application.services -> application.audit
application.audit -> application.database
""",
        [
            ["F1", "Browser/API client", "FastAPI routes", "JSON, form data, query params, headers", "Request size limit, Pydantic validation, neutral errors."],
            ["F2", "Static frontend /app", "FastAPI routes", "same-origin API calls", "No token persistence; API data rendered with textContent."],
            ["F3", "FastAPI routes", "Dependencies", "Bearer token and DB session", "JWT validation, active user check, session lifecycle in dependency."],
            ["F4", "Dependencies", "Authorization matrix", "current user role and object relation", "Role checks plus object-level status transition checks."],
            ["F5", "Services", "SQLAlchemy ORM / DB", "validated business objects", "Expression API, bounded pagination, no raw SQL string construction."],
            ["F6", "Services", "Audit log helper", "structured event details", "No password/token logging; controlled action names."],
        ],
    )

    add_heading(doc, "Сверка требований P5.")
    add_table(
        doc,
        ["Требование P5", "Что выполнено в проекте", "Доказательство"],
        [
            [
                "Задание 1: память, ресурсы, сериализация, курсоры БД, критичный сценарий",
                "Разобран сценарий unbounded offset как resource exhaustion для DB-backed списков и отчетов; добавлен settings.list_max_offset и Query(le=...).",
                "tests/test_security_hardening.py::test_huge_pagination_offsets_are_rejected; pytest 24 passed.",
            ],
            [
                "Задание 2: secure code review входов и минимум 3 sink",
                "Построена карта source-propagation-sink-protection: SQLAlchemy DB, audit JSON, report JSON, refresh token store, frontend DOM.",
                "D2-код trust boundaries, таблица потоков, таблица sink и подтверждение защит в коде.",
            ],
            [
                "Задание 3: authentication, authorization, cryptography",
                "Исправлены object-level authorization, user directory exposure, login timing discrepancy и production secret policy.",
                "Новые tests/test_auth.py, tests/test_users_directory.py, tests/test_maintenance_requests.py, tests/test_security_hardening.py.",
            ],
            [
                "Задание 4: 5 новых уязвимостей, CVSS v4.0 и CWE",
                "Выбраны P5-01..P5-05, рассчитаны CVSS v4.0 Base vectors, указаны CWE и приоритет исправления.",
                "Таблица CVSS в разделе 4; официальный FIRST CVSS v4.0 calculator.",
            ],
            [
                "Отчет: схема, flowchart, vulnerability table, role/access matrix, SAST/SCA, logs",
                "Все перечисленные элементы включены в отчет; диаграммы представлены D2-кодом и таблицами потоков.",
                "Разделы 1-5 данного отчета.",
            ],
        ],
        [5.6, 12.3, 8.4],
        font_size=9.2,
    )


def add_task1(doc: Document) -> None:
    add_heading(doc, "Задание 1. Аудит памяти, ресурсов и обработки данных.")
    add_paragraph(
        doc,
        "В MVP нет загрузки файлов и ручного управления памятью, поэтому основной практический риск "
        "находится в обработке больших входных данных и списочных запросов к базе. Лекция L7 "
        "рассматривает опасную обработку данных как ситуацию, где пользовательский ввод меняет "
        "поведение программы или приводит к DoS. В нашем случае attacker-controlled offset не "
        "выполняется как код, но может заставить БД выполнять дорогую операцию пропуска большого "
        "числа строк.",
    )
    add_table(
        doc,
        ["Участок", "Риск до P5", "Защита после P5", "Остаточный риск"],
        [
            [
                "HTTP request body",
                "Слишком большое тело запроса могло расходовать память до бизнес-логики.",
                "В P4 уже есть middleware с request size limit и обработкой 413.",
                "Низкий: файловых upload endpoints нет.",
            ],
            [
                "List/report query offset",
                "offset имел только ge=0, поэтому огромные значения доходили до SQL offset.",
                "Добавлен settings.list_max_offset=10000 и le=... во всех list/report маршрутах.",
                "Низкий: для больших production datasets лучше перейти на cursor pagination.",
            ],
            [
                "DB sessions",
                "Риск утечки соединений при исключениях.",
                "get_db_session использует generator dependency и закрывает Session в finally.",
                "Низкий при штатном FastAPI lifecycle.",
            ],
            [
                "JSON serialization",
                "Риск утечки чувствительных полей при прямом возврате моделей.",
                "Ответы идут через Pydantic response_model; refresh token хранится только как hash.",
                "Низкий; audit details отдельно проверяются на отсутствие секретов.",
            ],
        ],
        [5.0, 7.5, 8.2, 5.6],
    )
    add_paragraph(doc, "Критичный сценарий P5-03: authenticated supervisor отправляет запрос с offset=1000000000 к спискам активов, заявок, пользователей или отчету. До исправления API принимал такой offset и передавал его в SQLAlchemy .offset(offset). После исправления FastAPI отклоняет offset выше configured maximum с HTTP 422.")
    add_paragraph(doc, "Фрагмент кода до исправления (P5-03):")
    add_code(
        doc,
        annotated_snippet(
            [
                "Задание 1 / Таблица рисков: строка 'List/report query offset'.",
                "До P5: параметр offset имеет только нижнюю границу ge=0.",
                "Связь с P5-03: source = query parameter, sink = SQLAlchemy .offset(offset).",
            ],
            "materials/p4-snapshots/mvp-after/app/api/routes/assets.py",
            45,
            52,
        ),
    )
    add_paragraph(doc, "Фрагмент кода после исправления (P5-03):")
    add_code(
        doc,
        annotated_snippet(
            [
                "Задание 1 / Таблица рисков: строка 'List/report query offset'.",
                "После P5: добавлена верхняя граница le=settings.list_max_offset.",
                "Доказательство: test_huge_pagination_offsets_are_rejected ожидает HTTP 422.",
            ],
            "mvp/app/api/routes/assets.py",
            45,
            52,
        ),
    )
    add_paragraph(doc, "То же ограничение добавлено в routes/maintenance_requests.py, routes/users.py и routes/reports.py. Проверка: tests/test_security_hardening.py::test_huge_pagination_offsets_are_rejected возвращает 422 для всех четырех endpoint.")


def add_task2(doc: Document) -> None:
    add_heading(doc, "Задание 2. Secure code review входов данных и dangerous sinks.")
    add_paragraph(
        doc,
        "Для основного бизнес-сценария 'создание и выполнение заявки на обслуживание' построена цепочка "
        "source -> propagation -> sink -> protection. В MVP нет файловой системы как пользовательского "
        "sink и нет внешних сервисов, поэтому они отмечены как not applicable; вместо этого проверены "
        "реально существующие sinks: SQLAlchemy DB, refresh token store, audit JSON, report JSON и DOM-rendering.",
    )
    add_d2_listing(
        doc,
        "Листинг 2. D2-код source -> propagation -> sink -> protection.",
        "Код фиксирует основные точки входа пользовательских данных, распространение внутри MVP, dangerous sinks и меры защиты. D2-источник может быть отрендерен в SVG и использован как схема в draw.io или Word.",
        """
direction: right

sources: {
  label: "Sources"
  login: "Login form: email/password"
  token: "Bearer token"
  params: "Path and query params"
  body: "JSON request body"
  filters: "Report filters"
}

propagation: {
  label: "Propagation"
  routes: "FastAPI routes"
  schemas: "Pydantic schemas"
  dependencies: "Dependencies"
  services: "Service layer"
  authz: "Authorization checks"
}

sinks: {
  label: "Dangerous sinks"
  auth: "Password verification and token issuance"
  db: "SQLAlchemy DB queries"
  status_update: "Maintenance status update"
  audit_json: "Audit JSON details"
  report_json: "Report JSON response"
  dom: "Frontend DOM"
}

protections: {
  label: "Protections"
  argon2: "Argon2 and dummy hash"
  orm_bounds: "ORM expressions and bounded pagination"
  object_matrix: "Role plus object matrix"
  no_secret_logging: "No secret logging"
  response_models: "Pydantic response models"
  text_content: "textContent rendering"
}

sources.login -> propagation.routes: "email/password"
propagation.routes -> sinks.auth: "credential verification"
sinks.auth -> protections.argon2
sources.token -> propagation.dependencies: "JWT"
propagation.dependencies -> propagation.authz
propagation.authz -> sinks.status_update
sinks.status_update -> protections.object_matrix
sources.params -> propagation.schemas
propagation.schemas -> propagation.services
propagation.services -> sinks.db
sinks.db -> protections.orm_bounds
sources.body -> propagation.schemas: "status payload"
sources.filters -> propagation.routes
propagation.services -> sinks.report_json
sinks.report_json -> protections.response_models
sinks.audit_json -> protections.no_secret_logging
sinks.report_json -> sinks.dom
sinks.dom -> protections.text_content
""",
        [
            ["F1", "Login form", "Password verification/token sink", "email and password", "Argon2 verification, dummy hash for absent/inactive/locked user, neutral error."],
            ["F2", "Bearer token", "Authorization decision", "JWT claims and current user", "JWT validation, active user lookup, role/object matrix."],
            ["F3", "Query params", "SQLAlchemy DB query", "limit, offset, role/status filters", "Typed enums, ge/le bounds, ORM expression API."],
            ["F4", "JSON request body", "Maintenance status update", "status, assigned_engineer_id, internal_notes", "Pydantic validator and backend authorization check."],
            ["F5", "Request context", "Audit JSON details", "actor, action, entity, outcome", "Structured details without raw passwords or tokens."],
            ["F6", "Report filters", "Report JSON and frontend DOM", "summary rows and status counts", "Role-protected report endpoint, response_model, textContent rendering."],
        ],
    )
    add_table(
        doc,
        ["Source", "Propagation", "Sink", "Protection"],
        [
            [
                "OAuth2 login form: username/password",
                "routes/auth.py -> authenticate_user() -> password verification",
                "Password hash check and token issuance",
                "Argon2/password context, neutral error, dummy hash for absent/inactive/locked user.",
            ],
            [
                "Bearer token",
                "api/deps.py -> get_current_user() -> DB lookup",
                "Protected route authorization decision",
                "JWT validation, active user check, role and object checks in authorization.py.",
            ],
            [
                "Path parameter request_id",
                "FastAPI UUID parsing -> service lookup",
                "MaintenanceRequest row update",
                "UUID type validation, object existence check, status transition matrix.",
            ],
            [
                "Query params limit/offset/status",
                "Pydantic/FastAPI Query -> service list functions",
                "SQLAlchemy SELECT with offset/limit",
                "Typed enums, ge/le bounds, SQLAlchemy expression API, no string SQL construction.",
            ],
            [
                "Audit metadata from request/user/action",
                "AuditEvent -> create_audit_log()",
                "audit_logs.details JSON",
                "Structured dictionary, no raw passwords/tokens, controlled event names.",
            ],
            [
                "Report query filters",
                "routes/reports.py -> build_maintenance_summary()",
                "Report JSON response",
                "Role-filtered data access, bounded pagination, response_model serialization.",
            ],
            [
                "API JSON displayed in browser",
                "frontend/app.js API calls -> UI rendering",
                "DOM",
                "Use text rendering rather than concatenated executable HTML for data values.",
            ],
        ],
        [5.3, 7.0, 6.2, 8.0],
        font_size=8.9,
    )
    add_paragraph(doc, "Проверка SQL sink: в проекте используются SQLAlchemy select()/where()/offset()/limit(), а не ручная конкатенация SQL. Это снижает риск SQL injection, описанный в L7 как смешивание данных и команды.")
    add_code(
        doc,
        annotated_snippet(
            [
                "Задание 2 / Таблица source-propagation-sink: строка 'Query params limit/offset/status'.",
                "Propagation 4: FastAPI Query передает limit/offset/status в service layer.",
                "Sink 1: SQLAlchemy SELECT; защита = expression API без string SQL concatenation.",
            ],
            "mvp/app/services/users.py",
            12,
            21,
        ),
    )
    add_paragraph(doc, "Проверка report JSON sink: отчет строится из ORM objects в заранее заданную Pydantic schema, а не из произвольной сериализации пользовательского объекта.")
    add_code(
        doc,
        annotated_snippet(
            [
                "Задание 2 / Таблица source-propagation-sink: строка 'Report query filters'.",
                "Propagation 6: route -> build_maintenance_summary() -> response_model.",
                "Sink 4: Report JSON; защита = role-filtered data access и bounded pagination.",
            ],
            "mvp/app/api/routes/reports.py",
            18,
            39,
        ),
    )


def add_task3(doc: Document) -> None:
    add_heading(doc, "Задание 3. Authentication, authorization и cryptography.")
    add_paragraph(
        doc,
        "L8 разделяет факт входа в систему и право на конкретное действие. В P5 именно это "
        "разделение было усилено: supervisor остается привилегированной ролью для назначения и "
        "обзора, но не получает право выполнять работу вместо назначенного engineer. Также устранены "
        "side-channel timing discrepancy на login path и риск production bootstrap credentials.",
    )
    add_d2_listing(
        doc,
        "Листинг 3. D2-код блок-схемы authorization decision.",
        "Код ниже отражает два сценария из задания 3: легитимное выполнение работ назначенным engineer и запрещенную попытку supervisor выполнить работу вместо engineer.",
        """
direction: down

login: "1. Login request"
authn: "2. Authentication: active user and password hash"
token: "3. Access token issued"
load: "4. Load current user and request object"
assign: "5. Supervisor assigns engineer"
engineer_check: {
  label: "6. Actor is assigned engineer?"
  shape: diamond
}
start: "7. Move assigned request to in_progress"
complete: "8. Move assigned request to completed"
audit: "9. Audit event and report data"
deny: {
  label: "403 Forbidden"
  shape: hexagon
}
directory_check: {
  label: "Directory access requested?"
  shape: diamond
}
engineer_directory: "Engineer-only directory"
full_directory: "Full role directory"

login -> authn
authn -> token
token -> load
load -> assign
assign -> engineer_check
engineer_check -> start: "yes"
start -> complete
complete -> audit
engineer_check -> deny: "no: supervisor/admin attempts start or complete"
deny -> audit
load -> directory_check
directory_check -> engineer_directory: "supervisor"
directory_check -> full_directory: "technical_admin"
""",
        [
            ["F1", "Login request", "Authentication service", "email and password", "Password hash verification, dummy hash path, neutral error."],
            ["F2", "Access token", "Current user dependency", "Bearer token", "JWT validation and active user lookup."],
            ["F3", "Supervisor", "Assigned request", "assign engineer action", "Privileged assign/cancel allowed by role check."],
            ["F4", "Assigned engineer", "Status transition", "in_progress and completed", "Object-level check: actor id equals assigned_engineer_id."],
            ["F5", "Supervisor/admin", "Forbidden status transition", "start/complete attempt", "HTTP 403 because privileged role is not assigned engineer."],
            ["F6", "Directory request", "User listing response", "role filter", "Supervisor forced to engineer scope; technical admin can list all roles."],
        ],
    )
    add_table(
        doc,
        ["Действие", "Engineer", "Supervisor", "Technical admin", "Контроль"],
        [
            ["Login", "Да", "Да", "Да", "Authentication: active user + password hash."],
            ["List own/assigned maintenance requests", "Да", "Да", "Да", "Role-filtered service logic."],
            ["Create maintenance request", "Да", "Да", "Да", "Authenticated user, Pydantic schema."],
            ["Assign engineer", "Нет", "Да", "Да", "Privileged role check."],
            ["Start assigned work", "Только назначенный", "Нет", "Нет", "Object-level check in authorization.py."],
            ["Complete assigned work", "Только назначенный", "Нет", "Нет", "Object-level check in authorization.py."],
            ["List users", "Только себя", "Только engineers", "Все роли", "Directory scope in users.py."],
            ["Maintenance summary report", "Свои данные", "Обзор", "Обзор", "Role-aware report query."],
        ],
        [5.1, 3.4, 3.8, 3.8, 10.4],
        font_size=8.8,
    )
    add_paragraph(doc, "P5-01. До исправления supervisor/admin проходили early return в ensure_can_transition_request() и могли переводить заявку в in_progress/completed.")
    add_paragraph(doc, "Фрагмент кода до исправления (P5-01):")
    add_code(
        doc,
        annotated_snippet(
            [
                "Задание 3 / Role matrix: строки 'Start assigned work' и 'Complete assigned work'.",
                "До P5: authenticated supervisor/admin получает early return без object-level authorization.",
                "Задание 4 / P5-01 / CWE-863 / CVSS 7.1 High.",
            ],
            "materials/p4-snapshots/mvp-after/app/core/authorization.py",
            52,
            59,
        ),
    )
    add_paragraph(doc, "После исправления privileged roles могут назначать/cancel, но start/complete остается только назначенному engineer.")
    add_paragraph(doc, "Фрагмент кода после исправления (P5-01):")
    add_code(
        doc,
        annotated_snippet(
            [
                "Задание 3 / Role matrix: supervisor может assign/cancel, но не start/complete.",
                "После P5: object-level check отделяет право входа от права выполнить действие.",
                "Доказательство: test_supervisor_cannot_start_or_complete_assigned_engineer_work.",
            ],
            "mvp/app/core/authorization.py",
            52,
            74,
        ),
    )
    add_paragraph(doc, "P5-02. До исправления supervisor мог запрашивать всех пользователей или technical_admin через directory endpoint.")
    add_paragraph(doc, "Фрагмент кода до исправления (P5-02):")
    add_code(
        doc,
        annotated_snippet(
            [
                "Задание 3 / Role matrix: строка 'List users'.",
                "До P5: supervisor мог перечислять роли шире engineer directory.",
                "Задание 4 / P5-02 / CWE-200 / CVSS 5.3 Medium.",
            ],
            "materials/p4-snapshots/mvp-after/app/api/routes/users.py",
            17,
            34,
        ),
    )
    add_paragraph(doc, "После исправления supervisor directory принудительно ограничен engineers, а broad directory оставлен только technical admin.")
    add_paragraph(doc, "Фрагмент кода после исправления (P5-02):")
    add_code(
        doc,
        annotated_snippet(
            [
                "Задание 3 / Role matrix: supervisor видит только engineers.",
                "После P5: role filter None или ENGINEER разрешен, другие роли получают HTTP 403.",
                "Доказательство: test_supervisor_directory_is_limited_to_engineers.",
            ],
            "mvp/app/api/routes/users.py",
            17,
            43,
        ),
    )
    add_paragraph(doc, "P5-04. Login timing discrepancy устранен dummy hash verification: отсутствующий, inactive или locked user теперь проходит сравнимый password verification path перед нейтральной ошибкой.")
    add_paragraph(doc, "Фрагмент кода до исправления (P5-04):")
    add_code(
        doc,
        annotated_snippet(
            [
                "Задание 3 / Authentication risk: missing или inactive user возвращался без password verification.",
                "Задание 4 / P5-04 / CWE-203 / CVSS 6.9 Medium.",
                "Риск: ветки missing-user и real-user могли отличаться по времени выполнения.",
            ],
            "materials/p4-snapshots/mvp-after/app/services/auth.py",
            58,
            70,
        ),
    )
    add_paragraph(doc, "Фрагмент кода после исправления (P5-04):")
    add_code(
        doc,
        annotated_multirange_snippet(
            [
                "Задание 3 / Authentication risk: login path не должен раскрывать existence через timing.",
                "Задание 4 / P5-04 / CWE-203 / CVSS 6.9 Medium.",
                "Protection: dummy Argon2 hash делает missing/inactive/locked branch похожим на real-user branch.",
            ],
            "mvp/app/services/auth.py",
            [(27, 27), (63, 76)],
        ),
    )
    add_paragraph(doc, "P5-05. Production secret policy дополнена запретом DEBUG=true и placeholder bootstrap passwords.")
    add_paragraph(doc, "Фрагмент кода до исправления (P5-05):")
    add_code(
        doc,
        annotated_snippet(
            [
                "Задание 3 / Cryptography and secrets: проверялся SECRET_KEY, но не DEBUG и bootstrap passwords.",
                "Задание 4 / P5-05 / CWE-798 / CVSS 9.2 Critical.",
                "Риск: production мог стартовать с placeholder bootstrap credentials.",
            ],
            "materials/p4-snapshots/mvp-after/app/core/config.py",
            68,
            77,
        ),
    )
    add_paragraph(doc, "Фрагмент кода после исправления (P5-05):")
    add_code(
        doc,
        annotated_snippet(
            [
                "Задание 3 / Cryptography and secrets: production не должен принимать placeholder credentials.",
                "Задание 4 / P5-05 / CWE-798 / CVSS 9.2 Critical.",
                "Protection: DEBUG=false и реальные bootstrap passwords обязательны для production.",
            ],
            "mvp/app/core/config.py",
            80,
            100,
        ),
    )
    add_table(
        doc,
        ["Сценарий", "Ожидаемое решение", "Фактическое подтверждение"],
        [
            [
                "Легитимный: assigned engineer переводит назначенную заявку в in_progress и completed.",
                "Разрешить, потому что role=engineer и object ownership совпадает с assigned_engineer_id.",
                "tests/test_maintenance_requests.py покрывает allowed engineer status transitions.",
            ],
            [
                "Запрещенный: supervisor пытается start/complete работу назначенного engineer.",
                "Запретить HTTP 403, потому что supervisor не является исполнителем работ.",
                "test_supervisor_cannot_start_or_complete_assigned_engineer_work.",
            ],
            [
                "Запрещенный: supervisor запрашивает /users?role=technical_admin.",
                "Запретить HTTP 403 и не раскрывать admin directory.",
                "test_supervisor_directory_is_limited_to_engineers.",
            ],
        ],
        [8.6, 8.6, 9.1],
        font_size=9.2,
    )


def add_task4(doc: Document) -> None:
    add_heading(doc, "Задание 4. Пять уязвимостей, CWE и CVSS v4.0.")
    add_paragraph(
        doc,
        "Оценки рассчитаны как CVSS v4.0 Base vectors по официальному калькулятору FIRST. "
        "Приоритет исправления учитывает не только численный балл, но и бизнес-риски oil and gas "
        "maintenance MVP: целостность статуса работ, доступ к привилегированным аккаунтам и "
        "устойчивость API под нагрузкой.",
    )
    add_table(
        doc,
        ["ID", "Уязвимость", "Модуль/endpoint", "CWE", "CVSS v4.0 vector", "Score", "Priority"],
        [
            [
                "P5-05",
                "Predictable bootstrap passwords accepted in production",
                ".env.example / Settings",
                "CWE-798",
                "CVSS:4.0/AV:N/AC:L/AT:P/PR:N/UI:N/VC:H/VI:H/VA:L/SC:N/SI:N/SA:N",
                "9.2 Critical",
                "1",
            ],
            [
                "P5-01",
                "Supervisor could start/complete assigned engineer work",
                "PATCH /maintenance-requests/{id}/status",
                "CWE-863",
                "CVSS:4.0/AV:N/AC:L/AT:N/PR:L/UI:N/VC:N/VI:H/VA:N/SC:N/SI:L/SA:N",
                "7.1 High",
                "2",
            ],
            [
                "P5-04",
                "Login timing reveals account existence",
                "POST /auth/token",
                "CWE-203",
                "CVSS:4.0/AV:N/AC:L/AT:N/PR:N/UI:N/VC:L/VI:N/VA:N/SC:N/SI:N/SA:N",
                "6.9 Medium",
                "3",
            ],
            [
                "P5-02",
                "Supervisor could enumerate admins/all users",
                "GET /users",
                "CWE-200",
                "CVSS:4.0/AV:N/AC:L/AT:N/PR:L/UI:N/VC:L/VI:N/VA:N/SC:N/SI:N/SA:N",
                "5.3 Medium",
                "4",
            ],
            [
                "P5-03",
                "Huge pagination offsets reach DB queries",
                "GET list/report routes",
                "CWE-400",
                "CVSS:4.0/AV:N/AC:L/AT:N/PR:L/UI:N/VC:N/VI:N/VA:L/SC:N/SI:N/SA:N",
                "5.3 Medium",
                "5",
            ],
        ],
        [2.0, 4.9, 4.3, 2.4, 8.6, 2.5, 1.8],
        font_size=7.3,
    )
    add_table(
        doc,
        ["ID", "Source -> sink", "Impact", "Fix", "Proof fixed"],
        [
            [
                "P5-05",
                "Production env/config -> privileged bootstrap login",
                "Initial admin/supervisor/engineer account compromise.",
                "Reject DEBUG=true and placeholder bootstrap passwords in production.",
                "test_production_rejects_debug_and_placeholder_bootstrap_passwords.",
            ],
            [
                "P5-01",
                "Supervisor token -> status update sink",
                "Integrity loss: work can be started/completed by non-assigned user.",
                "State transition matrix requires assigned engineer for in_progress/completed.",
                "test_supervisor_cannot_start_or_complete_assigned_engineer_work.",
            ],
            [
                "P5-04",
                "Login form -> password verification branch",
                "Account enumeration support for password spraying.",
                "Dummy Argon2 hash verification on absent/inactive/locked user.",
                "test_missing_user_login_still_runs_password_verification.",
            ],
            [
                "P5-02",
                "Supervisor token -> user directory response",
                "Internal admin/identity exposure.",
                "Supervisor directory scope forced to engineers only.",
                "test_supervisor_directory_is_limited_to_engineers.",
            ],
            [
                "P5-03",
                "Query offset -> SQL offset",
                "Resource consumption and degraded availability.",
                "Configured maximum offset at API boundary.",
                "test_huge_pagination_offsets_are_rejected.",
            ],
        ],
        [2.0, 5.2, 6.2, 7.0, 6.1],
        font_size=8.7,
    )
    add_paragraph(
        doc,
        "Вывод по CVSS и бизнес-риску: самый высокий CVSS у P5-05, и бизнес-приоритет также самый "
        "высокий, потому что production bootstrap password может сразу открыть privileged access. "
        "P5-01 имеет меньший CVSS, но для предметной области он почти так же важен: ложное выполнение "
        "работ по нефтегазовым активам влияет на operational integrity. P5-03 имеет medium score, "
        "однако в production с крупной БД его риск растет, поэтому в дальнейшем рекомендуется cursor pagination.",
    )


def add_verification(doc: Document) -> None:
    add_heading(doc, "Инструменты, снапшоты и доказательства исправлений.")
    add_table(
        doc,
        ["Проверка", "Результат 2026-05-04", "Комментарий"],
        [
            ["Snapshot integrity", "Текущий mvp отличается от P4-after только ожидаемыми P5 source/test/config файлами.", "Кодовая база не разъехалась случайно."],
            ["pytest", "24 passed", "Функциональные и security regression tests."],
            ["ruff", "All checks passed", "Формат/линтинг без нарушений."],
            ["mypy", "Success: no issues found in 50 source files", "Типизация app и tests."],
            ["bandit", "No issues identified", "Python SAST security rules."],
            ["semgrep", "0 findings; 456 rules; 45 targets", "SAST-проверка по правилам Semgrep."],
            ["pip-audit", "No known vulnerabilities found", "p3-universal-mvp skipped because local package is not on PyPI."],
            ["pylint", "8.13/10", "Остались style warnings и известные false positives для Pydantic/SQLAlchemy."],
        ],
        [5.1, 8.0, 13.2],
        font_size=9.0,
    )
    add_paragraph(doc, "Ключевые команды валидации:")
    add_code(
        doc,
        """
.\\.venv\\Scripts\\python.exe -m pytest
.\\.venv\\Scripts\\python.exe -m ruff check .
.\\.venv\\Scripts\\python.exe -m mypy app tests
.\\.venv\\Scripts\\python.exe -m bandit -c bandit.yaml -r app
PYTHONUTF8=1 semgrep scan --config auto --exclude __pycache__ app
.\\.venv\\Scripts\\python.exe -m pip_audit
.\\.venv\\Scripts\\python.exe -m pylint app tests --score=y --reports=n
""",
        size=10,
    )
    add_paragraph(doc, "Использованные учебные и технические источники:")
    add_bullets(
        doc,
        [
            "Практическая работа №5: Comprehensive MVP security analysis.",
            "Лекция L7: Malware Injection and Data Processing.",
            "Лекция L8: Authentication, Authorization, Cryptography.",
            "FIRST CVSS v4.0 specification: https://www.first.org/cvss/v4.0/specification-document",
            "FIRST CVSS calculator: https://www.first.org/cvss/calculator/v4-0",
        ],
    )


def add_conclusion(doc: Document) -> None:
    add_heading(doc, "Заключение.")
    add_paragraph(
        doc,
        "Практическая работа №5 выполнена как продолжение P4: исходный P4-after снапшот был принят "
        "как baseline, после чего проведен ручной и автоматизированный анализ по требованиям P5. "
        "В коде исправлены пять новых уязвимостей: excessive authorization для supervisor, user "
        "directory overexposure, unbounded pagination offset, login timing discrepancy и production "
        "acceptance of predictable bootstrap credentials.",
    )
    add_paragraph(
        doc,
        "Результаты подтверждены regression tests, SAST, SCA и статической проверкой типов. В отчет "
        "включены D2-код структурной схемы, карта trust boundaries, flowchart принятия решения "
        "о доступе, role/access matrix, таблица уязвимостей, классификация CWE, CVSS v4.0 vectors, "
        "фрагменты кода до и после исправления, а также proof-fixed tests. Для LMS вместе с отчетом "
        "предоставляется папка mvp с исправленным исходным кодом.",
    )
    add_paragraph(
        doc,
        "Подтверждающие материалы представлены в виде воспроизводимых команд валидации и таблиц "
        "результатов. Формат соответствует требованию P5 о наличии screenshots or logs proving fixes: "
        "для каждого исправления указан проверочный тест или инструментальный результат, а диаграммы "
        "даны как D2-код, пригодный для рендеринга в SVG и последующего импорта в draw.io или Word.",
    )


def build_report() -> Path:
    doc = Document()
    configure_document(doc)
    add_title(doc)
    add_intro(doc)
    add_task1(doc)
    add_task2(doc)
    add_task3(doc)
    add_task4(doc)
    add_verification(doc)
    add_conclusion(doc)
    return save_with_fallback(doc, REPORT_PATH)


if __name__ == "__main__":
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    path = build_report()
    print(path)

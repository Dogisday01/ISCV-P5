from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
INPUT_PATH = ROOT / "docs" / "p3-report-draft.md"
OUTPUT_PATH = ROOT / "output" / "doc" / "p3-report-draft.docx"


def set_run_font(run, size: int, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def configure_styles(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.79)
    section.bottom_margin = Inches(0.79)
    section.left_margin = Inches(1.18)
    section.right_margin = Inches(0.59)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor(0, 0, 0)

    heading_1 = document.styles["Heading 1"]
    heading_1.font.name = "Times New Roman"
    heading_1._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    heading_1.font.size = Pt(16)
    heading_1.font.bold = True
    heading_1.font.color.rgb = RGBColor(0, 0, 0)

    heading_2 = document.styles["Heading 2"]
    heading_2.font.name = "Times New Roman"
    heading_2._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    heading_2.font.size = Pt(14)
    heading_2.font.bold = True
    heading_2.font.color.rgb = RGBColor(0, 0, 0)

    heading_3 = document.styles["Heading 3"]
    heading_3.font.name = "Times New Roman"
    heading_3._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    heading_3.font.size = Pt(12)
    heading_3.font.bold = True
    heading_3.font.color.rgb = RGBColor(0, 0, 0)


def add_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(text)
    set_run_font(run, 12)


def add_code_block(document: Document, text: str, language: str | None) -> None:
    if language:
        label = document.add_paragraph()
        label.alignment = WD_ALIGN_PARAGRAPH.LEFT
        label_run = label.add_run(f"РљРѕРґ {language}:")
        set_run_font(label_run, 12, bold=True)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Inches(0.25)
    run = paragraph.add_run(text.rstrip())
    set_run_font(run, 12)


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for row_idx, row in enumerate(rows):
        for col_idx, cell_text in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(cell_text)
            set_run_font(run, 12, bold=row_idx == 0)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].startswith("|"):
        raw_cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        if all(set(cell) <= {"-", ":"} for cell in raw_cells):
            index += 1
            continue
        rows.append(raw_cells)
        index += 1
    return rows, index


def build_document() -> None:
    document = Document()
    configure_styles(document)

    lines = INPUT_PATH.read_text(encoding="utf-8").splitlines()
    index = 0

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if not stripped:
            index += 1
            continue

        if stripped.startswith("```"):
            language = stripped[3:].strip() or None
            index += 1
            block_lines: list[str] = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                block_lines.append(lines[index])
                index += 1
            add_code_block(document, "\n".join(block_lines), language)
            index += 1
            continue

        if stripped.startswith("|"):
            rows, index = parse_table(lines, index)
            add_table(document, rows)
            continue

        if stripped.startswith("# "):
            paragraph = document.add_paragraph(style="Heading 1")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(stripped[2:].strip())
            set_run_font(run, 16, bold=True)
            index += 1
            continue

        if stripped.startswith("## "):
            paragraph = document.add_paragraph(style="Heading 2")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(stripped[3:].strip())
            set_run_font(run, 14, bold=True)
            index += 1
            continue

        if stripped.startswith("### "):
            paragraph = document.add_paragraph(style="Heading 3")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(stripped[4:].strip())
            set_run_font(run, 12, bold=True)
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            next_line = lines[index].strip()
            if not next_line or next_line.startswith("#") or next_line.startswith("|") or next_line.startswith("```"):
                break
            paragraph_lines.append(next_line)
            index += 1
        add_paragraph(document, " ".join(paragraph_lines))

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
